import io
import re
import logging
from typing import Dict, Any

logger = logging.getLogger("firstfintech")


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text content from an uploaded file.
    Supports: .txt, .pdf, .docx, .md
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "txt":
        return _extract_txt(file_bytes)
    elif ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext == "docx":
        return _extract_docx(file_bytes)
    elif ext == "md":
        return _extract_markdown(file_bytes)
    else:
        logger.warning(f"Unsupported file type: .{ext}")
        return f"[Unsupported file type: .{ext}]"


def extract_text_with_metadata(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract text and return metadata about the document."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    text = extract_text_from_file(file_bytes, filename)

    words = len(text.split()) if text else 0
    lines = text.count("\n") + 1 if text else 0
    sections = len(re.findall(r"(?m)^#{1,3}\s|^[A-Z][A-Z\s]{3,}$|^\d+\.\s+[A-Z]", text)) if text else 0

    meta = {
        "text": text,
        "filename": filename,
        "file_type": ext,
        "word_count": words,
        "line_count": lines,
        "section_count": sections,
        "page_count": 0,
    }

    if ext == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            meta["page_count"] = len(reader.pages)
        except Exception:
            pass

    return meta


def _extract_txt(file_bytes: bytes) -> str:
    """Decode raw bytes as UTF-8 text."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        return "\n".join(pages_text)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return f"[Error extracting PDF: {e}]"


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX including tables and headings."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            # Preserve heading hierarchy
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level_num = int(level)
                except ValueError:
                    level_num = 1
                parts.append(f"{'#' * level_num} {para.text}")
            else:
                parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append("\n[TABLE]\n" + "\n".join(rows_text) + "\n[/TABLE]")

        return "\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return f"[Error extracting DOCX: {e}]"


def _extract_markdown(file_bytes: bytes) -> str:
    """Extract text from a Markdown file."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")
    return text
